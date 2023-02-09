from docx2pdf import convert
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import datetime


def __main__():
    # try:
        count_files = len(sys.argv)-1

        input_date = input("YY/MM/DD: ")
        date = datetime.datetime.strptime(input_date, "%Y/%m/%d")

        for i in range(count_files):
            append_text = "{}/{}/{}".format(date.year, date.month, date.day)
            file_path = sys.argv[i+1]

            doc = docx.Document(file_path)
            print("docx file '{}' found")

            doc.add_paragraph(append_text)
            para = doc.paragraphs[-1]
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.runs[0].font
            doc.save(file_path)

            convert(file_path)

            date += datetime.timedelta(days=1)

    # except Exception as e:
    #     print(e)
    #     input()


__main__()
